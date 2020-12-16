import os
import re
import cv2
import tempfile
import unicodedata
import pytesseract
import preprocessing
import matplotlib.pyplot as plt
from docx import Document
from pdf2image import convert_from_path
from pytesseract import Output

# This is a functional version of the OCR program

IMG_EXT = ["bmp", "pnm", "png", "jfif", "jpeg", "jpg", "tiff"]


def pdf_to_img(pdf_file):
    return convert_from_path(pdf_file)


def ocr_core(img_file, preprocesses):
    if preprocesses:
        img_file = preprocessing.preprocess(img_file, preprocesses)
    return pytesseract.image_to_string(img_file)


def check_ext(file):
    return file.split(".")[-1]


def ocr_core_pdf(pdf_file, preprocesses):
    images = pdf_to_img(pdf_file)
    if len(images) > 1:
        text_pages = []
        for pg, image in enumerate(images):
            temp = tempfile.gettempdir()
            img_path = os.path.join(temp, "temp.jpg")
            image.save(img_path, "JPEG")
            text_pages.append(ocr_core(img_path, preprocesses))
        return text_pages
    else:
        return ocr_core(images[0], preprocesses)


def orientation_script_detection(img_file):
    image = cv2.imread(img_file)
    osd = pytesseract.image_to_osd(image)
    angle = re.search(r"(?<=Rotate: )\d+", osd).group(0)
    script = re.search(r"(?<=Script: )\w+", osd).group(0)
    return angle, script


def remove_control_characters(s):
    return "".join(ch for ch in s if unicodedata.category(ch)[0] != "C")


def save_file(output, file_type="txt", save_path=".", save_file_name="output"):
    if not os.path.exists(save_path):
        print("Output path doesn't exist. Output file will be saved in current directory")
        save_path = ".."
    if file_type == "docx":
        document = Document()
        if isinstance(output, list):
            for page in output:
                page = remove_control_characters(page)
                document.add_paragraph(page)
                document.add_page_break()
        else:
            output = remove_control_characters(output)
            document.add_paragraph(output)
        document.save(save_path + f"/{save_file_name}.docx")
    elif file_type == "txt":
        with open(os.path.join(save_path + f"/{save_file_name}.txt"), "w") as temp_file:
            if isinstance(output, list):
                for page in output:
                    temp_file.write(page)
            else:
                temp_file.write(output)


def pdf_to_text(file, preprocesses=None, save=False, file_type="txt", save_path=".", save_file_name="pdf_output"):
    ext = check_ext(file)
    if ext == "pdf":
        output = ocr_core_pdf(file, preprocesses)
        if save:
            save_file(output, file_type, save_path, save_file_name)
        else:
            if isinstance(output, list):
                for page in output:
                    print(page)
            else:
                print(output)
    else:
        print("Please provide pdf file only. If you want to convert image file to text, use 'img_to_txt' function.")


def img_to_text(file, preprocesses=None, save=False, file_type="txt", save_path=".", save_file_name="img_output"):
    ext = check_ext(file)
    if ext in IMG_EXT:
        output = ocr_core(file, preprocesses)
        if save:
            save_file(output, file_type, save_path, save_file_name)
        else:
            print(output)
    else:
        print(f"Please provide image file with {IMG_EXT} formats only!")


def show_image(image, fig_size, title):
    b, g, r = cv2.split(image)
    rgb_img = cv2.merge([r, g, b])
    plt.figure(figsize=fig_size)
    plt.imshow(rgb_img)
    plt.title(title)
    plt.show()


def character_bounding_box(img_file, fig_size=(8, 6), box_color=(0, 255, 0), box_width=2, title="Character Bounded Box"):
    image = cv2.imread(img_file)
    h, w, c = image.shape
    boxes = pytesseract.image_to_boxes(image)
    for b in boxes.splitlines():
        b = b.split(" ")
        image = cv2.rectangle(image, (int(b[1]), h - int(b[2])), (int(b[3]), h - int(b[4])), box_color, box_width)
    show_image(image, fig_size, title)


def word_bounding_box(img_file, confidence=60, fig_size=(8, 6), box_color=(0, 255, 0), box_width=2, title="Word Bounded Box"):
    image = cv2.imread(img_file)
    d = pytesseract.image_to_data(image, output_type=Output.DICT)
    # print("DATA KEYS: \n", d.keys())
    n_boxes = len(d["text"])
    for i in range(n_boxes):
        # condition to only pick boxes with a confidence > 60%
        if int(d["conf"][i]) > confidence:
            (x, y, w, h) = (d["left"][i], d["top"][i], d["width"][i], d["height"][i])
            image = cv2.rectangle(image, (x, y), (x + w, y + h), box_color, box_width)
    show_image(image, fig_size, title)


def pattern_bounding_box(img_file, reg_pat, confidence=60, fig_size=(8, 6), box_color=(0, 255, 0), box_width=2, title="Pattern Bounded Box"):
    image = cv2.imread(img_file)
    d = pytesseract.image_to_data(image, output_type=Output.DICT)
    n_boxes = len(d["text"])
    for i in range(n_boxes):
        if int(d["conf"][i]) > confidence:
            if re.match(reg_pat, d["text"][i]):
                (x, y, w, h) = (d["left"][i], d["top"][i], d["width"][i], d["height"][i])
                image = cv2.rectangle(image, (x, y), (x + w, y + h), box_color, box_width)
    show_image(image, fig_size, title)


print(orientation_script_detection("../images/sample2.png"))
