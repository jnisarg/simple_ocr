import os
import re
import cv2
import tempfile
import unicodedata
import pytesseract
import preprocessing
from docx import Document
import matplotlib.pyplot as plt
from pdf2image import convert_from_path

plt.style.use("ggplot")
plt.figure(figsize=(8, 8))

# This is the main OCR file which does all the working behind converting images/pdfs to text or drawing bounded boxes
# around the characters/words/pattern found in the file. This all works on the top of pytesseract module (python
# version of tesseract)


def preprocesses_info():
    """
    Displays available preprocesses
    """
    preprocesses_list = [preprocess.__name__ for preprocess in preprocessing.preprocesses_list]
    print(f"Preprocesses available are: {preprocesses_list}")


def pdf_to_img(pdf_file):
    """
    Converts pdf to images
    :param pdf_file: pdf file
    :returns: list of images
    """
    return convert_from_path(pdf_file)


def check_ext(file):
    """
    :param file: file for which extension to be checked
    :returns: extension of the given file
    """
    return file.split(".")[-1]


def remove_control_characters(s):
    """
    :param s: text
    :returns: text without any null or escape characters
    """
    return "".join(ch for ch in s if unicodedata.category(ch)[0] != "C")


def save_file(output, file_type, save_path, save_file_name):
    """
    Saves the output of extracted text from images/pdf to the specified file_type.
    :param output: output text
    :param file_type: txt or docx
    :param save_path:  location at which output file is to be saved
    :param save_file_name: name of the output file
    """
    if not os.path.exists(save_path):
        print("Output path doesn't exist. Output file will be saved in current directory")
        save_path = "."
    if file_type == "docx":
        document = Document()
        if isinstance(output, list):
            for page in output:
                page = remove_control_characters(page)
                document.add_paragraph(page)
                document.add_page_break()
        else:
            output = remove_control_characters(output)
            document.add_paragraph(output)
        document.save(os.path.join(save_path, f"{save_file_name}.docx"))
    elif file_type == "txt":
        with open(os.path.join(save_path + f"/{save_file_name}.txt"), "w") as tmp_file:
            if isinstance(output, list):
                for page in output:
                    tmp_file.write(page)
            else:
                tmp_file.write(output)


def show_image(image, fig_size, title):
    """
    Displays image
    :param image: image file
    :param fig_size: matplotlib figure size
    :param title: title of the figure
    """
    b, g, r = cv2.split(image)
    rgb_img = cv2.merge([r, g, b])
    plt.figure(figsize=fig_size)
    plt.imshow(rgb_img)
    plt.title(title)
    plt.show()


def ocr_core(img_file, preprocesses, config):
    """
    :param img_file: image file
    :param preprocesses: list of preprocesses to be done on image
    :param config: tesseract's config file
    :returns: text extracted from image
    """
    if preprocesses:
        img_file = preprocessing.preprocess(img_file, preprocesses)
    return pytesseract.image_to_string(img_file, config=config)


class OCR:

    IMG_EXT = ["bmp", "pnm", "png", "jfif", "jpeg", "jpg", "tiff", "webp", "ppm"]

    def __init__(self, file, config):
        """
        :param file: image/pdf file
        :param config: config for tesseract
        """
        self._file = file
        self._config = config
        self.input_type = check_ext(self._file)
        self.delete = False
        if self.input_type == "pdf":
            self.images = pdf_to_img(self._file)
        elif self.input_type in OCR.IMG_EXT:
            self.image = self._file
        else:
            print(f"Input file is not supported. Supported file types are: {['pdf'] + OCR.IMG_EXT}")
            self.delete = True

    def ocr_core_pdf(self, preprocesses):
        """
        Converts pdf to images and then returns text extracted from each image.
        :param preprocesses: list of preprocesses to be done on image
        :returns: text extracted from pdf file
        """
        if len(self.images) > 1:
            text_pages = []
            for pg, image in enumerate(self.images):
                temp = tempfile.gettempdir()
                img_path = os.path.join(temp, "temp.jpg")
                image.save(img_path, "JPEG")
                text_pages.append(ocr_core(img_path, preprocesses, self._config))
            return text_pages
        else:
            return ocr_core(self.images[0], preprocesses, self._config)

    def orientation_script_detection(self):
        """
        Detects orientation and script of the file
        :returns: angle and script
        """
        image = cv2.imread(self.image)
        osd = pytesseract.image_to_osd(image)
        angle = re.search(r"(?<=Rotate: )\d+", osd).group(0)
        script = re.search(r"(?<=Script: )\w+", osd).group(0)
        return angle, script

    def pdf_to_text(self, preprocesses=None, save=False, file_type="txt", save_path=".", save_file_name="pdf_output"):
        """
        Extracts text from pdf file
        :param preprocesses: list of preprocesses to be done on image
        :param save: bool
        :param file_type: txt or docx default: txt
        :param save_path: location at which output file is to be saved
        :param save_file_name: name of the output file
        """
        output = self.ocr_core_pdf(preprocesses)
        if save:
            save_file(output, file_type, save_path, save_file_name)
        else:
            if isinstance(output, list):
                for page in output:
                    print(page)
            else:
                print(output)

    def img_to_text(self, preprocesses=None, save=False, file_type="txt", save_path=".", save_file_name="img_output"):
        """
        Extracts text from image
        :param preprocesses: list of preprocesses to be done on image
        :param save: bool
        :param file_type: txt or docx default: txt
        :param save_path: location at which output file is to be saved
        :param save_file_name: name of the output file
        """
        output = ocr_core(self.image, preprocesses, self._config)
        if save:
            save_file(output, file_type, save_path, save_file_name)
        else:
            print(output)

    def character_bounding_box(self, fig_size=(8, 6), box_color=(0, 255, 0), box_width=2, title="Character Bounded Boxes"):
        """
        Displays boxes around each character in the file
        :param fig_size: figure size
        :param box_color: (r, g, b)
        :param box_width: int width of the box
        :param title: title of the figure
        """
        image = cv2.imread(self.image)
        h, w, c = image.shape
        boxes = pytesseract.image_to_boxes(image)
        for b in boxes.splitlines():
            b = b.split(" ")
            image = cv2.rectangle(image, (int(b[1]), h - int(b[2])), (int(b[3]), h - int(b[4])), box_color, box_width)
        show_image(image, fig_size, title)

    def word_bounding_box(self, confidence=60, fig_size=(8, 6), box_color=(0, 255, 0), box_width=2, title="Word Bounded Boxes"):
        """
        Displays boxes around each word in the file
        :param confidence: confidence of the word in the file
        :param fig_size: figure size
        :param box_color: (r, g, b)
        :param box_width: int width of the box
        :param title: title of the figure
        """
        image = cv2.imread(self.image)
        d = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
        n_boxes = len(d["text"])
        for i in range(n_boxes):
            if int(d["conf"][i]) > confidence:
                (x, y, w, h) = (d["left"][i], d["top"][i], d["width"][i], d["height"][i])
                image = cv2.rectangle(image, (x, y), (x + w, y + h), box_color, box_width)
        show_image(image, fig_size, title)

    def pattern_bounding_box(self, reg_pat, confidence=60, fig_size=(8, 6), box_color=(0, 255, 0), box_width=2, title="Pattern Bounded Boxes"):
        """
        Displays boxes around matched patten in the file
        :param reg_pat: pattern to be matched
        :param confidence: confidence of the word in the file
        :param fig_size: figure size
        :param box_color: (r, g, b)
        :param box_width: int width of the box
        :param title: title of the figure
        """
        image = cv2.imread(self.image)
        d = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
        n_boxes = len(d["text"])
        for i in range(n_boxes):
            if int(d["conf"][i]) > confidence:
                if re.match(reg_pat, d["text"][i]):
                    (x, y, w, h) = (d["left"][i], d["top"][i], d["width"][i], d["height"][i])
                    image = cv2.rectangle(image, (x, y), (x + w, y + h), box_color, box_width)
        show_image(image, fig_size, title)
